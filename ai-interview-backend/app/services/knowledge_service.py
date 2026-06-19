import os
import logging
from typing import List, Optional, Dict

from fastapi import UploadFile
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, func, delete

from app.models.knowledge_document import KnowledgeDocument, KnowledgeChunk
from app.schemas.knowledge import (
    KnowledgeDocumentResponse, KnowledgeChunkResponse,
    KnowledgeRetrieveResult, KnowledgeTestRetrieve
)
from app.services.embedding_service import embedding_service
from app.exceptions.http_exceptions import APIException

logger = logging.getLogger(__name__)


class KnowledgeService:

    # ------------------------------------------------------------------ #
    # Text extraction helpers
    # ------------------------------------------------------------------ #

    @staticmethod
    def _extract_text(file_path: str, doc_type: str) -> str:
        """从文件中提取纯文本内容"""
        try:
            if doc_type == "pdf":
                import pdfplumber
                texts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            texts.append(page_text)
                return "\n".join(texts)

            elif doc_type == "docx":
                from docx import Document
                doc = Document(file_path)
                return "\n".join(p.text for p in doc.paragraphs)

            elif doc_type in ("txt", "md"):
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()

            else:
                raise APIException(
                    code=10001,
                    message=f"不支持的文档类型: {doc_type}",
                    status_code=400
                )
        except APIException:
            raise
        except Exception as e:
            logger.error(f"提取文本失败: {e}")
            raise APIException(
                code=10002,
                message=f"提取文本失败: {str(e)}",
                status_code=500
            )

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """将文本按指定大小和重叠度切分为多个块"""
        chunks: List[str] = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            start += chunk_size - overlap
        return chunks

    # ------------------------------------------------------------------ #
    # Mapping helpers
    # ------------------------------------------------------------------ #

    @staticmethod
    def _doc_to_response(doc: KnowledgeDocument) -> KnowledgeDocumentResponse:
        return KnowledgeDocumentResponse(
            id=doc.id,
            title=doc.title,
            filename=doc.filename,
            doc_type=doc.doc_type,
            category=doc.category,
            description=doc.description,
            file_size=doc.file_size,
            chunk_count=doc.chunk_count,
            is_vectorized=doc.is_vectorized,
            is_enabled=doc.is_enabled,
            created_at=doc.created_at,
            updated_at=doc.updated_at,
        )

    @staticmethod
    def _chunk_to_response(chunk: KnowledgeChunk) -> KnowledgeChunkResponse:
        return KnowledgeChunkResponse(
            id=chunk.id,
            document_id=chunk.document_id,
            chunk_index=chunk.chunk_index,
            content=chunk.content,
            is_vectorized=chunk.is_vectorized,
            created_at=chunk.created_at,
            updated_at=chunk.updated_at,
        )

    # ------------------------------------------------------------------ #
    # Document CRUD
    # ------------------------------------------------------------------ #

    @staticmethod
    async def list_documents(
        db: AsyncSession,
        page: int = 1,
        per_page: int = 20,
        keyword: Optional[str] = None,
        category: Optional[str] = None,
        is_enabled: Optional[bool] = None,
    ) -> Dict:
        """分页查询文档列表"""
        query = select(KnowledgeDocument)
        count_query = select(func.count(KnowledgeDocument.id))

        if keyword:
            query = query.where(KnowledgeDocument.title.ilike(f"%{keyword}%"))
            count_query = count_query.where(KnowledgeDocument.title.ilike(f"%{keyword}%"))
        if category is not None:
            query = query.where(KnowledgeDocument.category == category)
            count_query = count_query.where(KnowledgeDocument.category == category)
        if is_enabled is not None:
            query = query.where(KnowledgeDocument.is_enabled == is_enabled)
            count_query = count_query.where(KnowledgeDocument.is_enabled == is_enabled)

        # 总数
        total_result = await db.execute(count_query)
        total = total_result.scalar() or 0

        # 分页
        query = query.order_by(KnowledgeDocument.id.desc())
        query = query.offset((page - 1) * per_page).limit(per_page)
        result = await db.execute(query)
        docs = result.scalars().all()

        return {
            "items": [KnowledgeService._doc_to_response(d) for d in docs],
            "total": total,
            "per_page": per_page,
            "current_page": page,
        }

    @staticmethod
    async def get_document(db: AsyncSession, id: int) -> KnowledgeDocument:
        """获取文档 ORM 对象，不存在则抛 404"""
        result = await db.execute(
            select(KnowledgeDocument).where(KnowledgeDocument.id == id)
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise APIException(code=10004, message="文档不存在", status_code=404)
        return doc

    @staticmethod
    async def get_document_response(db: AsyncSession, id: int) -> KnowledgeDocumentResponse:
        """获取文档响应体"""
        doc = await KnowledgeService.get_document(db, id)
        return KnowledgeService._doc_to_response(doc)

    @staticmethod
    async def upload_document(
        db: AsyncSession,
        file: UploadFile,
        title: Optional[str] = None,
        category: Optional[str] = None,
        description: Optional[str] = None,
    ) -> KnowledgeDocumentResponse:
        """上传文档并执行分块 + 向量化"""
        os.makedirs("uploads/knowledge", exist_ok=True)
        filename = file.filename or "unknown"
        file_path = os.path.join("uploads/knowledge", filename)

        # 保存文件
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)

        # 推断文档类型
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        doc_type = ext if ext in ("pdf", "docx", "txt", "md") else "txt"

        # 标题为空时使用文件名（去掉扩展名）
        doc_title = title.strip() if title and title.strip() else filename.rsplit(".", 1)[0]

        doc = KnowledgeDocument(
            title=doc_title,
            filename=filename,
            doc_type=doc_type,
            category=category,
            description=description,
            file_path=file_path,
            file_size=len(content),
        )
        db.add(doc)
        await db.commit()
        await db.refresh(doc)

        # 提取、分块、向量化
        await KnowledgeService._process_document(db, doc)

        return KnowledgeService._doc_to_response(doc)

    @staticmethod
    async def _process_document(db: AsyncSession, doc: KnowledgeDocument):
        """提取文本 -> 分块 -> 写入数据库 -> 生成向量"""
        text = KnowledgeService._extract_text(doc.file_path, doc.doc_type)
        chunks_text = KnowledgeService._chunk_text(text)

        # 创建 chunk 记录
        chunk_objects: List[KnowledgeChunk] = []
        for idx, chunk_content in enumerate(chunks_text):
            chunk = KnowledgeChunk(
                document_id=doc.id,
                chunk_index=idx,
                content=chunk_content,
            )
            db.add(chunk)
            chunk_objects.append(chunk)

        doc.chunk_count = len(chunk_objects)
        await db.flush()

        # 批量生成 embeddings
        if chunk_objects:
            embeddings = await embedding_service.get_embeddings_batch(
                [c.content for c in chunk_objects]
            )
            for chunk, emb in zip(chunk_objects, embeddings):
                if emb is not None:
                    chunk.embedding = emb
                    chunk.is_vectorized = True

        # 如果有任意 chunk 已向量化，标记文档为已向量化
        if any(c.is_vectorized for c in chunk_objects):
            doc.is_vectorized = True

        await db.commit()
        await db.refresh(doc)

    @staticmethod
    async def delete_document(db: AsyncSession, id: int):
        """删除文档（级联删除 chunks）"""
        doc = await KnowledgeService.get_document(db, id)
        await db.delete(doc)
        await db.commit()

    @staticmethod
    async def toggle_document(
        db: AsyncSession, id: int, is_enabled: bool
    ) -> KnowledgeDocumentResponse:
        """启用 / 禁用文档"""
        doc = await KnowledgeService.get_document(db, id)
        doc.is_enabled = is_enabled
        await db.commit()
        await db.refresh(doc)
        return KnowledgeService._doc_to_response(doc)

    @staticmethod
    async def reindex_document(db: AsyncSession, id: int) -> KnowledgeDocumentResponse:
        """重新索引：删除旧分块后重新提取、分块、向量化"""
        doc = await KnowledgeService.get_document(db, id)

        # 删除旧 chunks
        await db.execute(
            delete(KnowledgeChunk).where(KnowledgeChunk.document_id == doc.id)
        )
        doc.chunk_count = 0
        doc.is_vectorized = False
        await db.flush()

        # 重新处理
        await KnowledgeService._process_document(db, doc)

        return KnowledgeService._doc_to_response(doc)

    # ------------------------------------------------------------------ #
    # Chunks
    # ------------------------------------------------------------------ #

    @staticmethod
    async def list_chunks(
        db: AsyncSession,
        document_id: int,
        page: int = 1,
        per_page: int = 20,
    ) -> Dict:
        """分页查询指定文档的分块列表"""
        query = select(KnowledgeChunk).where(KnowledgeChunk.document_id == document_id)
        count_query = select(func.count(KnowledgeChunk.id)).where(
            KnowledgeChunk.document_id == document_id
        )

        total_result = await db.execute(count_query)
        total = total_result.scalar() or 0

        query = query.order_by(KnowledgeChunk.chunk_index)
        query = query.offset((page - 1) * per_page).limit(per_page)
        result = await db.execute(query)
        chunks = result.scalars().all()

        return {
            "items": [KnowledgeService._chunk_to_response(c) for c in chunks],
            "total": total,
            "per_page": per_page,
            "current_page": page,
        }

    # ------------------------------------------------------------------ #
    # Retrieval
    # ------------------------------------------------------------------ #

    @staticmethod
    async def test_retrieve(
        db: AsyncSession, data: KnowledgeTestRetrieve
    ) -> List[KnowledgeRetrieveResult]:
        """语义检索测试"""
        # 生成查询向量
        query_embedding = await embedding_service.get_embedding(data.query)
        if query_embedding is None:
            raise APIException(
                code=10003,
                message="生成查询向量失败，请检查 embedding 服务配置",
                status_code=500,
            )

        # 构建过滤条件
        filters = [KnowledgeChunk.is_vectorized == True]  # noqa: E712

        if data.category:
            # 需要关联文档表过滤分类
            doc_ids_query = select(KnowledgeDocument.id).where(
                KnowledgeDocument.category == data.category,
                KnowledgeDocument.is_enabled == True,  # noqa: E712
            )
            result = await db.execute(doc_ids_query)
            doc_ids = [row[0] for row in result.all()]
            filters.append(KnowledgeChunk.document_id.in_(doc_ids))
        else:
            # 默认只检索已启用文档的 chunks
            enabled_doc_ids_query = select(KnowledgeDocument.id).where(
                KnowledgeDocument.is_enabled == True,  # noqa: E712
            )
            result = await db.execute(enabled_doc_ids_query)
            doc_ids = [row[0] for row in result.all()]
            filters.append(KnowledgeChunk.document_id.in_(doc_ids))

        # 向量检索
        search_results = await embedding_service.vector_search(
            db, KnowledgeChunk, query_embedding, filters, data.top_k
        )

        # 组装结果
        results: List[KnowledgeRetrieveResult] = []
        for item in search_results:
            chunk: KnowledgeChunk = item["instance"]
            # 加载父文档获取标题
            doc_result = await db.execute(
                select(KnowledgeDocument).where(
                    KnowledgeDocument.id == chunk.document_id
                )
            )
            doc = doc_result.scalar_one_or_none()

            results.append(
                KnowledgeRetrieveResult(
                    similarity_score=item["similarity_score"],
                    chunk_content=chunk.content,
                    document_id=chunk.document_id,
                    document_title=doc.title if doc else "未知文档",
                    chunk_index=chunk.chunk_index,
                )
            )

        return results


knowledge_service = KnowledgeService()
